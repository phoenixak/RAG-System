"""
Document Processing System
Handles text extraction from multiple file formats with error recovery.
"""

import csv
from abc import ABC, abstractmethod
from typing import Any, Dict, List, Optional, Tuple

import chardet
import pandas as pd
import pdfplumber
from docx import Document
from PyPDF2 import PdfReader

from src.core.config import get_settings
from src.core.logging import LoggerMixin
from src.documents.models import DocumentType

settings = get_settings()


class DocumentProcessor(ABC, LoggerMixin):
    """Base class for document processors."""

    @abstractmethod
    def can_process(self, file_path: str, content_type: str) -> bool:
        """Check if this processor can handle the given file."""
        pass

    @abstractmethod
    async def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from document.

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        pass

    @abstractmethod
    def get_document_type(self) -> DocumentType:
        """Get the document type this processor handles."""
        pass

    def detect_encoding(self, file_path: str, sample_size: int = 10000) -> str:
        """Detect file encoding using chardet."""
        try:
            with open(file_path, "rb") as f:
                raw_data = f.read(sample_size)

            result = chardet.detect(raw_data)
            encoding = result.get("encoding", "utf-8")
            confidence = result.get("confidence", 0)

            self.logger.debug(
                "Encoding detected",
                file_path=file_path,
                encoding=encoding,
                confidence=confidence,
            )

            # Fallback to utf-8 if confidence is too low
            if confidence < 0.7:
                encoding = "utf-8"

            return encoding

        except Exception as e:
            self.logger.warning(
                "Encoding detection failed, using utf-8",
                file_path=file_path,
                error=str(e),
            )
            return "utf-8"


class PDFProcessor(DocumentProcessor):
    """PDF document processor using pdfplumber and PyPDF2."""

    def can_process(self, file_path: str, content_type: str) -> bool:
        """Check if file is a PDF."""
        return file_path.lower().endswith(".pdf") or content_type in ["application/pdf"]

    def get_document_type(self) -> DocumentType:
        """Return PDF document type."""
        return DocumentType.PDF

    async def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using pdfplumber with PyPDF2 fallback."""
        try:
            # Try pdfplumber first (better text extraction)
            return await self._extract_with_pdfplumber(file_path)
        except Exception as e:
            self.logger.warning(
                "pdfplumber extraction failed, trying PyPDF2",
                file_path=file_path,
                error=str(e),
            )
            try:
                return await self._extract_with_pypdf2(file_path)
            except Exception as e2:
                self.logger.error(
                    "Both PDF extraction methods failed",
                    file_path=file_path,
                    pdfplumber_error=str(e),
                    pypdf2_error=str(e2),
                )
                raise

    async def _extract_with_pdfplumber(
        self, file_path: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text using pdfplumber."""
        text_parts = []
        metadata = {
            "extractor": "pdfplumber",
            "page_content": {},
            "tables": [],
            "images": [],
        }

        with pdfplumber.open(file_path) as pdf:
            # Extract document metadata
            pdf_metadata = pdf.metadata or {}
            metadata.update(
                {
                    "title": pdf_metadata.get("Title"),
                    "author": pdf_metadata.get("Author"),
                    "subject": pdf_metadata.get("Subject"),
                    "creator": pdf_metadata.get("Creator"),
                    "producer": pdf_metadata.get("Producer"),
                    "creation_date": pdf_metadata.get("CreationDate"),
                    "modification_date": pdf_metadata.get("ModDate"),
                    "page_count": len(pdf.pages),
                }
            )

            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                page_text = page.extract_text() or ""

                # Extract tables
                tables = page.extract_tables()
                if tables:
                    table_text = self._format_tables(tables)
                    page_text += "\n" + table_text
                    metadata["tables"].extend(
                        [{"page": page_num, "table_count": len(tables)}]
                    )

                # Store page content mapping
                if page_text.strip():
                    metadata["page_content"][page_num] = {
                        "start_char": len("\n".join(text_parts)),
                        "end_char": len("\n".join(text_parts)) + len(page_text),
                        "text_length": len(page_text),
                    }
                    text_parts.append(page_text)

        full_text = "\n".join(text_parts)
        metadata["text_length"] = len(full_text)

        self.logger.info(
            "PDF text extracted with pdfplumber",
            file_path=file_path,
            page_count=metadata["page_count"],
            text_length=len(full_text),
            table_count=len(metadata["tables"]),
        )

        return full_text, metadata

    async def _extract_with_pypdf2(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text using PyPDF2 as fallback."""
        text_parts = []
        metadata = {"extractor": "PyPDF2", "page_content": {}}

        with open(file_path, "rb") as file:
            pdf_reader = PdfReader(file)

            # Extract document metadata
            pdf_metadata = pdf_reader.metadata or {}
            metadata.update(
                {
                    "title": pdf_metadata.get("/Title"),
                    "author": pdf_metadata.get("/Author"),
                    "subject": pdf_metadata.get("/Subject"),
                    "creator": pdf_metadata.get("/Creator"),
                    "producer": pdf_metadata.get("/Producer"),
                    "creation_date": pdf_metadata.get("/CreationDate"),
                    "modification_date": pdf_metadata.get("/ModDate"),
                    "page_count": len(pdf_reader.pages),
                }
            )

            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text() or ""

                    if page_text.strip():
                        metadata["page_content"][page_num] = {
                            "start_char": len("\n".join(text_parts)),
                            "end_char": len("\n".join(text_parts)) + len(page_text),
                            "text_length": len(page_text),
                        }
                        text_parts.append(page_text)

                except Exception as e:
                    self.logger.warning(
                        "Failed to extract text from page",
                        file_path=file_path,
                        page_num=page_num,
                        error=str(e),
                    )

        full_text = "\n".join(text_parts)
        metadata["text_length"] = len(full_text)

        self.logger.info(
            "PDF text extracted with PyPDF2",
            file_path=file_path,
            page_count=metadata["page_count"],
            text_length=len(full_text),
        )

        return full_text, metadata

    def _format_tables(self, tables: List[List[List[str]]]) -> str:
        """Format extracted tables as text."""
        formatted_tables = []

        for table in tables:
            if not table:
                continue

            # Convert table to pandas DataFrame for better formatting
            try:
                df = pd.DataFrame(table[1:], columns=table[0] if table else [])
                table_text = df.to_string(index=False, na_rep="")
                formatted_tables.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
            except Exception:
                # Fallback to simple formatting
                table_rows = []
                for row in table:
                    if row:
                        table_rows.append(" | ".join(str(cell) for cell in row if cell))
                if table_rows:
                    formatted_tables.append(
                        "\n[TABLE]\n" + "\n".join(table_rows) + "\n[/TABLE]\n"
                    )

        return "\n".join(formatted_tables)


class DOCXProcessor(DocumentProcessor):
    """DOCX document processor using python-docx."""

    def can_process(self, file_path: str, content_type: str) -> bool:
        """Check if file is a DOCX."""
        return file_path.lower().endswith(".docx") or content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]

    def get_document_type(self) -> DocumentType:
        """Return DOCX document type."""
        return DocumentType.DOCX

    async def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX document."""
        try:
            doc = Document(file_path)

            text_parts = []
            metadata = {
                "extractor": "python-docx",
                "paragraph_count": 0,
                "table_count": 0,
                "image_count": 0,
                "style_info": {},
            }

            # Extract document properties
            props = doc.core_properties
            metadata.update(
                {
                    "title": props.title,
                    "author": props.author,
                    "subject": props.subject,
                    "keywords": props.keywords,
                    "category": props.category,
                    "comments": props.comments,
                    "created": props.created,
                    "modified": props.modified,
                    "last_modified_by": props.last_modified_by,
                }
            )

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # Track styles
                    style = para.style.name if para.style else "Normal"
                    metadata["style_info"][style] = (
                        metadata["style_info"].get(style, 0) + 1
                    )

                    # Add style information as markup for important headings
                    if "Heading" in style:
                        text_parts.append(f"\n# {para.text.strip()}\n")
                    else:
                        text_parts.append(para.text.strip())

                    metadata["paragraph_count"] += 1

            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
                    metadata["table_count"] += 1

            # Count images/shapes
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    metadata["image_count"] += 1

            full_text = "\n".join(text_parts)
            metadata["text_length"] = len(full_text)

            self.logger.info(
                "DOCX text extracted",
                file_path=file_path,
                paragraph_count=metadata["paragraph_count"],
                table_count=metadata["table_count"],
                image_count=metadata["image_count"],
                text_length=len(full_text),
            )

            return full_text, metadata

        except Exception as e:
            self.logger.error(
                "DOCX extraction failed", file_path=file_path, error=str(e)
            )
            raise

    def _extract_table_text(self, table) -> str:
        """Extract text from a Word table."""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells.append(cell_text)
            if any(cells):  # Only add non-empty rows
                rows.append(" | ".join(cells))

        return "\n".join(rows)


class TXTProcessor(DocumentProcessor):
    """Plain text document processor."""

    def can_process(self, file_path: str, content_type: str) -> bool:
        """Check if file is plain text."""
        return file_path.lower().endswith(".txt") or content_type.startswith(
            "text/plain"
        )

    def get_document_type(self) -> DocumentType:
        """Return TXT document type."""
        return DocumentType.TXT

    async def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file."""
        encoding = self.detect_encoding(file_path)

        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as file:
                text = file.read()

            metadata = {
                "extractor": "text",
                "encoding": encoding,
                "line_count": text.count("\n") + 1,
                "text_length": len(text),
            }

            self.logger.info(
                "Text file processed",
                file_path=file_path,
                encoding=encoding,
                line_count=metadata["line_count"],
                text_length=len(text),
            )

            return text, metadata

        except Exception as e:
            self.logger.error(
                "Text extraction failed", file_path=file_path, error=str(e)
            )
            raise


class CSVProcessor(DocumentProcessor):
    """CSV document processor."""

    def can_process(self, file_path: str, content_type: str) -> bool:
        """Check if file is CSV."""
        return file_path.lower().endswith(".csv") or content_type in [
            "text/csv",
            "application/csv",
        ]

    def get_document_type(self) -> DocumentType:
        """Return CSV document type."""
        return DocumentType.CSV

    async def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV file."""
        encoding = self.detect_encoding(file_path)

        try:
            # Try pandas first for better CSV handling
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                text = df.to_string(index=False)

                metadata = {
                    "extractor": "pandas",
                    "encoding": encoding,
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "columns": df.columns.tolist(),
                    "text_length": len(text),
                }

            except Exception:
                # Fallback to standard csv module
                with open(file_path, "r", encoding=encoding, errors="replace") as file:
                    csv_reader = csv.reader(file)
                    rows = list(csv_reader)

                if not rows:
                    return "", {
                        "extractor": "csv",
                        "encoding": encoding,
                        "row_count": 0,
                    }

                # Format as text
                text_rows = []
                headers = rows[0] if rows else []

                # Add headers
                if headers:
                    text_rows.append(" | ".join(headers))
                    text_rows.append("-" * len(" | ".join(headers)))

                # Add data rows
                for row in rows[1:]:
                    if row:
                        text_rows.append(" | ".join(str(cell) for cell in row))

                text = "\n".join(text_rows)
                metadata = {
                    "extractor": "csv",
                    "encoding": encoding,
                    "row_count": len(rows),
                    "column_count": len(headers) if headers else 0,
                    "columns": headers,
                    "text_length": len(text),
                }

            self.logger.info(
                "CSV file processed",
                file_path=file_path,
                encoding=encoding,
                row_count=metadata["row_count"],
                column_count=metadata["column_count"],
                text_length=len(text),
            )

            return text, metadata

        except Exception as e:
            self.logger.error(
                "CSV extraction failed", file_path=file_path, error=str(e)
            )
            raise


class DocumentProcessorFactory:
    """Factory for creating document processors."""

    def __init__(self):
        self.processors = [
            PDFProcessor(),
            DOCXProcessor(),
            TXTProcessor(),
            CSVProcessor(),
        ]

    def get_processor(
        self, file_path: str, content_type: str
    ) -> Optional[DocumentProcessor]:
        """Get appropriate processor for the file."""
        for processor in self.processors:
            if processor.can_process(file_path, content_type):
                return processor
        return None

    def get_supported_types(self) -> List[str]:
        """Get list of supported file extensions."""
        return [".pdf", ".docx", ".txt", ".csv"]

    def validate_file_type(self, file_path: str, content_type: str) -> bool:
        """Validate if file type is supported."""
        return self.get_processor(file_path, content_type) is not None


# Global factory instance
document_processor_factory = DocumentProcessorFactory()
